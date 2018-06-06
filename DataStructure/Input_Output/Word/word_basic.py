# -*- coding: utf-8 -*-
'''
和纯文本相比，.docx 文件有很多结构。这些结构在 python-docx 中用 3 种不同
的类型来表示。在最高一层，Document 对象表示整个文档。Document 对象包含
一个 Paragraph 对象的列表，表示文档中的段落（用户在 Word 文档中输入时，如
果按下回车，新的段落就开始了）。每个 Paragraph 对象都包含一个 Run 对象的列
表。
'''
import docx

#   读取 Word 文档
import docx


doc = docx.Document('demo.docx')
len(doc.paragraphs) # 7
print(doc.paragraphs[0].text) # 'Document Title'
print(doc.paragraphs[1].text)  # 'A plain paragraph with some bold and some italic'
len(doc.paragraphs[1].runs)   # 4
print(doc.paragraphs[1].runs[0].text) # 'A plain paragraph with some '
print(doc.paragraphs[1].runs[1].text) # 'bold'
print(doc.paragraphs[1].runs[2].text) # ' and some '
print(doc.paragraphs[1].runs[3].text) # 'italic'

#   从.docx 文件中取得完整的文本
#如果你只关心 Word 文档中的文本，不关心样式信息，就可以利用 getText()函数。
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n\n'.join(fullText)
print(getText('demo.docx'))

#   Run 属性
doc = docx.Document('demo.docx')
print(doc.paragraphs[0].text) # 'Document Title'
print(doc.paragraphs[0].style) # 'Title'
doc.paragraphs[0].style = 'Normal'
print(doc.paragraphs[1].text) # 'A plain paragraph with some bold and some italic'
print((doc.paragraphs[1].runs[0].text, doc.paragraphs[1].runs[1].text, doc.
paragraphs[1].runs[2].text, doc.paragraphs[1].runs[3].text)) # ('A plain paragraph with some ', 'bold', ' and some ', 'italic')
doc.paragraphs[1].runs[0].style = 'QuoteChar'
doc.paragraphs[1].runs[1].underline = True
doc.paragraphs[1].runs[3].underline = True
doc.save('restyled.docx')

#   写入 Word 文档
doc = docx.Document()
doc.add_paragraph('Hello world!') #   <docx.text.Paragraph object at 0x0000000003B56F60>
doc.add_paragraph('Hello world!', 'Title')
paraObj1 = doc.add_paragraph('This is a second paragraph.')
paraObj2 = doc.add_paragraph('This is a yet another paragraph.')
paraObj1.add_run(' This text is being added to the second paragraph.') # <docx.text.Run object at 0x0000000003A2C860>
doc.save('helloworld.docx')

#   添加标题
doc = docx.Document()
doc.add_heading('Header 0', 0)
doc.add_heading('Header 1', 1)
doc.add_heading('Header 2', 2)
doc.add_heading('Header 3', 3)
doc.save('headings.docx')

#   添加换行符和换页符
doc = docx.Document()
doc.add_paragraph('This is on the first page!')
doc.paragraphs[0].runs[0].add_break(docx.text.WD_BREAK.PAGE)
doc.add_paragraph('This is on the second page!')
doc.save('twoPage.docx')

#   添加图像
doc.add_picture('zophie.png', width=docx.shared.Inches(1),height=docx.shared.Cm(4))